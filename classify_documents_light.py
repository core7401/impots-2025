from __future__ import annotations

import csv
import json
import re
import shutil
import traceback
from pathlib import Path
from typing import Dict, List, Tuple

import requests
from docx import Document
from pypdf import PdfReader


CONFIG_FILE = Path("config_documents.json")


def load_config(config_path: Path) -> dict:
    print(f"[INFO] Chargement config : {config_path.resolve()}")

    if not config_path.exists():
        raise FileNotFoundError(f"Fichier de configuration introuvable : {config_path}")

    with config_path.open("r", encoding="utf-8") as f:
        config = json.load(f)

    required_keys = [
        "base_dir",
        "inbox_dir",
        "review_dir",
        "ollama",
        "supported_extensions",
        "thresholds",
        "categories",
        "apartments",
    ]

    for key in required_keys:
        if key not in config:
            raise ValueError(f"Clé manquante dans config : {key}")

    return config


def normalize_text(text: str) -> str:
    text = text.lower()
    replacements = {
        "é": "e", "è": "e", "ê": "e", "ë": "e",
        "à": "a", "â": "a", "ä": "a",
        "î": "i", "ï": "i",
        "ô": "o", "ö": "o",
        "ù": "u", "û": "u", "ü": "u",
        "ç": "c", "œ": "oe",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)

    text = text.replace("\n", " ").replace("\r", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_for_digits(text: str) -> str:
    return re.sub(r"\D", "", text)


def safe_filename(name: str) -> str:
    invalid = '<>:"/\\|?*'
    for ch in invalid:
        name = name.replace(ch, "_")
    return name.strip()


def sanitize_for_name(text: str) -> str:
    text = normalize_text(text)
    text = text.replace(" ", "_")
    text = re.sub(r"[^a-z0-9_]+", "_", text)
    text = re.sub(r"_+", "_", text)
    return text.strip("_")


def shorten_text(text: str, max_len: int) -> str:
    return text[:max_len]


def guess_year(text: str, fallback: str = "2025") -> str:
    matches = re.findall(r"\b(20\d{2})\b", text)
    if matches:
        return matches[-1]
    return fallback


class LightTextExtractor:
    def extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()

        if suffix in {".txt", ".md", ".csv"}:
            return self._read_text_file(file_path)

        if suffix == ".pdf":
            return self._read_pdf(file_path)

        if suffix == ".docx":
            return self._read_docx(file_path)

        raise RuntimeError(f"Format non pris en charge par la version légère : {suffix}")

    def _read_text_file(self, file_path: Path) -> str:
        encodings_to_try = ["utf-8", "utf-8-sig", "cp1252", "latin-1"]
        for enc in encodings_to_try:
            try:
                return file_path.read_text(encoding=enc, errors="strict")
            except Exception:
                pass

        return file_path.read_text(encoding="utf-8", errors="ignore")

    def _read_pdf(self, file_path: Path) -> str:
        reader = PdfReader(str(file_path))
        pages_text: List[str] = []

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
            except Exception as exc:
                print(f"[WARN] PDF page {i + 1} non lue dans {file_path.name}: {exc}")

        text = "\n".join(pages_text).strip()
        return text

    def _read_docx(self, file_path: Path) -> str:
        doc = Document(str(file_path))
        parts: List[str] = []

        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()


class OllamaClient:
    def __init__(self, base_url: str, model: str, timeout_seconds: int) -> None:
        self.base_url = base_url.rstrip("/")
        self.model = model
        self.timeout_seconds = timeout_seconds

    def classify_with_ai(
        self,
        filename: str,
        text: str,
        categories: List[str],
        apartments: List[str],
        apartment_hints: Dict[str, Dict[str, List[str]]],
    ) -> Tuple[str, str]:
        apartment_description_lines = []

        for apt_name, hints in apartment_hints.items():
            hints_flat = []
            for level in ("strong_keywords", "medium_keywords", "weak_keywords"):
                hints_flat.extend(hints.get(level, []))
            hints_preview = ", ".join(hints_flat[:8])
            apartment_description_lines.append(f"- {apt_name}: {hints_preview}")

        prompt = f"""
Tu classes un document immobilier.

Réponds en JSON strict :
{{
  "category": "...",
  "apartment": "..."
}}

Catégories autorisées :
{json.dumps(categories, ensure_ascii=False)}

Appartements autorisés :
{json.dumps(apartments, ensure_ascii=False)}

Indices possibles :
{chr(10).join(apartment_description_lines)}

Règles :
- Choisis exactement une catégorie autorisée.
- Choisis exactement un appartement autorisé.
- N'invente rien.
- Si le texte est partiel, prends l'option la plus probable parmi les choix autorisés.

Nom du fichier :
{filename}

Contenu :
{text}
""".strip()

        payload = {
            "model": self.model,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0}
        }

        response = requests.post(
            f"{self.base_url}/generate",
            json=payload,
            timeout=self.timeout_seconds
        )
        response.raise_for_status()
        data = response.json()

        raw = (data.get("response") or "").strip()

        try:
            start = raw.find("{")
            end = raw.rfind("}")
            if start != -1 and end != -1 and end > start:
                raw = raw[start:end + 1]

            parsed = json.loads(raw)
            category = str(parsed.get("category", "")).strip()
            apartment = str(parsed.get("apartment", "")).strip()

            if category not in categories:
                category = ""
            if apartment not in apartments:
                apartment = ""

            return category, apartment

        except Exception:
            return "", ""


def score_category(
    text_norm: str,
    filename_norm: str,
    categories: Dict[str, List[str]]
) -> Tuple[str, Dict[str, int]]:
    scores: Dict[str, int] = {}

    for category, keywords in categories.items():
        score = 0
        for kw in keywords:
            kw_norm = normalize_text(kw)
            if kw_norm in text_norm:
                score += 2
            if kw_norm in filename_norm:
                score += 3
        scores[category] = score

    best_category = max(scores, key=scores.get)
    return best_category, scores


def score_apartment(
    text_norm: str,
    filename_norm: str,
    apartments: Dict[str, Dict[str, List[str]]]
) -> Tuple[str, Dict[str, int]]:
    scores: Dict[str, int] = {}
    digits_text = normalize_for_digits(text_norm)
    digits_filename = normalize_for_digits(filename_norm)

    for apartment_name, criteria in apartments.items():
        score = 0

        for kw in criteria.get("strong_keywords", []):
            kw_norm = normalize_text(kw)
            kw_digits = normalize_for_digits(kw)

            if kw_norm and kw_norm in text_norm:
                score += 8
            if kw_norm and kw_norm in filename_norm:
                score += 10

            if len(kw_digits) >= 6:
                if kw_digits in digits_text:
                    score += 8
                if kw_digits in digits_filename:
                    score += 10

        for kw in criteria.get("medium_keywords", []):
            kw_norm = normalize_text(kw)
            kw_digits = normalize_for_digits(kw)

            if kw_norm and kw_norm in text_norm:
                score += 4
            if kw_norm and kw_norm in filename_norm:
                score += 6

            if len(kw_digits) >= 6:
                if kw_digits in digits_text:
                    score += 4
                if kw_digits in digits_filename:
                    score += 6

        for kw in criteria.get("weak_keywords", []):
            kw_norm = normalize_text(kw)
            if kw_norm and kw_norm in text_norm:
                score += 2
            if kw_norm and kw_norm in filename_norm:
                score += 3

        scores[apartment_name] = score

    best_apartment = max(scores, key=scores.get)
    return best_apartment, scores


def is_category_decision_confident(scores: Dict[str, int], min_category_score: int) -> bool:
    return bool(scores) and max(scores.values()) >= min_category_score


def is_apartment_decision_confident(
    scores: Dict[str, int],
    min_apartment_score: int,
    min_score_gap: int
) -> bool:
    if not scores:
        return False

    ordered = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    best_score = ordered[0][1]

    if best_score < min_apartment_score:
        return False

    if len(ordered) == 1:
        return True

    second_score = ordered[1][1]
    return (best_score - second_score) >= min_score_gap


def ensure_directories(
    inbox_dir: Path,
    review_dir: Path,
    base_dir: Path,
    categories: Dict[str, List[str]],
    apartments: Dict[str, Dict[str, List[str]]]
) -> None:
    print("[INFO] Vérification dossiers")
    inbox_dir.mkdir(parents=True, exist_ok=True)
    review_dir.mkdir(parents=True, exist_ok=True)

    for apartment in apartments.keys():
        for category in categories.keys():
            (base_dir / apartment / category).mkdir(parents=True, exist_ok=True)


def build_new_filename(category: str, apartment: str, year: str, suffix: str) -> str:
    category_part = sanitize_for_name(category)
    apartment_part = sanitize_for_name(apartment)
    return f"{category_part}_{apartment_part}_{year}{suffix.lower()}"


def move_to_destination(file_path: Path, apartment: str, category: str, base_dir: Path, year: str) -> Path:
    destination_dir = base_dir / apartment / category
    destination_dir.mkdir(parents=True, exist_ok=True)

    new_name = build_new_filename(category, apartment, year, file_path.suffix)
    destination = destination_dir / new_name

    counter = 1
    while destination.exists():
        destination = destination_dir / f"{destination.stem}_{counter}{file_path.suffix.lower()}"
        counter += 1

    shutil.move(str(file_path), str(destination))
    return destination


def move_to_review(file_path: Path, review_dir: Path, reason: str) -> Path:
    destination = review_dir / safe_filename(file_path.name)

    counter = 1
    while destination.exists():
        destination = review_dir / f"{file_path.stem}_{counter}{file_path.suffix}"
        counter += 1

    shutil.move(str(file_path), str(destination))

    log_file = review_dir / "review_log.txt"
    with log_file.open("a", encoding="utf-8") as f:
        f.write(f"{file_path.name} -> {reason}\n")

    return destination


def append_csv_log(
    base_dir: Path,
    original_name: str,
    final_path: str,
    apartment: str,
    category: str,
    year: str,
    status: str,
    category_scores: Dict[str, int],
    apartment_scores: Dict[str, int],
) -> None:
    log_path = base_dir / "classification_log.csv"
    file_exists = log_path.exists()

    with log_path.open("a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f, delimiter=";")
        if not file_exists:
            writer.writerow([
                "original_name",
                "final_path",
                "apartment",
                "category",
                "year",
                "status",
                "category_scores",
                "apartment_scores"
            ])

        writer.writerow([
            original_name,
            final_path,
            apartment,
            category,
            year,
            status,
            json.dumps(category_scores, ensure_ascii=False),
            json.dumps(apartment_scores, ensure_ascii=False),
        ])


def process_file(file_path: Path, extractor: LightTextExtractor, ollama: OllamaClient, config: dict) -> None:
    print(f"\n[INFO] Traitement : {file_path.name}")

    base_dir = Path(config["base_dir"])
    review_dir = Path(config["review_dir"])
    categories = config["categories"]
    apartments = config["apartments"]

    thresholds = config["thresholds"]
    min_apartment_score = int(thresholds["min_apartment_score"])
    min_score_gap = int(thresholds["min_score_gap"])
    min_category_score = int(thresholds["min_category_score"])
    max_text_for_ai = int(config["ollama"]["max_text_for_ai"])

    print("[INFO] Début extraction texte")
    raw_text = extractor.extract_text(file_path)
    print("[INFO] Fin extraction texte")

    text = raw_text.strip()
    print(f"[INFO] Longueur texte extrait : {len(text)}")

    if not text:
        destination = move_to_review(file_path, review_dir, "Aucun texte extrait")
        print(f"[WARN] Aucun texte extrait -> {destination}")
        append_csv_log(
            base_dir=base_dir,
            original_name=file_path.name,
            final_path=str(destination),
            apartment="",
            category="",
            year="",
            status="A_VERIFIER",
            category_scores={},
            apartment_scores={},
        )
        return

    filename_norm = normalize_text(file_path.name)
    text_norm = normalize_text(text)

    rule_category, category_scores = score_category(text_norm, filename_norm, categories)
    rule_apartment, apartment_scores = score_apartment(text_norm, filename_norm, apartments)

    category_confident = is_category_decision_confident(category_scores, min_category_score)
    apartment_confident = is_apartment_decision_confident(apartment_scores, min_apartment_score, min_score_gap)

    print(f"[DEBUG] Scores catégories : {category_scores}")
    print(f"[DEBUG] Scores appartements : {apartment_scores}")

    final_category = rule_category if category_confident else ""
    final_apartment = rule_apartment if apartment_confident else ""

    if not final_category or not final_apartment:
        print("[INFO] Règles insuffisantes -> appel Ollama")
        ai_category, ai_apartment = ollama.classify_with_ai(
            filename=file_path.name,
            text=shorten_text(text, max_text_for_ai),
            categories=list(categories.keys()),
            apartments=list(apartments.keys()),
            apartment_hints=apartments
        )
        print(f"[DEBUG] Réponse IA : catégorie={ai_category!r}, appartement={ai_apartment!r}")

        if not final_category and ai_category in categories:
            final_category = ai_category

        if not final_apartment and ai_apartment in apartments:
            final_apartment = ai_apartment

    if not final_category or not final_apartment:
        reason = f"Classification incomplète | category={final_category or 'None'} | apartment={final_apartment or 'None'}"
        destination = move_to_review(file_path, review_dir, reason)
        print(f"[WARN] Document ambigu -> {destination}")
        append_csv_log(
            base_dir=base_dir,
            original_name=file_path.name,
            final_path=str(destination),
            apartment=final_apartment,
            category=final_category,
            year="",
            status="A_VERIFIER",
            category_scores=category_scores,
            apartment_scores=apartment_scores,
        )
        return

    year = guess_year(text, "2025")
    destination = move_to_destination(file_path, final_apartment, final_category, base_dir, year)

    print(f"[OK] Appartement : {final_apartment}")
    print(f"[OK] Catégorie   : {final_category}")
    print(f"[OK] Année       : {year}")
    print(f"[OK] Destination : {destination}")

    append_csv_log(
        base_dir=base_dir,
        original_name=file_path.name,
        final_path=str(destination),
        apartment=final_apartment,
        category=final_category,
        year=year,
        status="OK",
        category_scores=category_scores,
        apartment_scores=apartment_scores,
    )


def main() -> None:
    print("[INFO] Démarrage script")

    config = load_config(CONFIG_FILE)

    base_dir = Path(config["base_dir"])
    inbox_dir = Path(config["inbox_dir"])
    review_dir = Path(config["review_dir"])
    categories = config["categories"]
    apartments = config["apartments"]
    supported_extensions = {ext.lower() for ext in config["supported_extensions"]}

    print(f"[INFO] base_dir   = {base_dir}")
    print(f"[INFO] inbox_dir  = {inbox_dir}")
    print(f"[INFO] review_dir = {review_dir}")

    ensure_directories(inbox_dir, review_dir, base_dir, categories, apartments)

    ollama = OllamaClient(
        base_url=config["ollama"]["base_url"],
        model=config["ollama"]["model"],
        timeout_seconds=int(config["ollama"]["timeout_seconds"])
    )
    extractor = LightTextExtractor()

    if not inbox_dir.exists():
        print(f"[WARN] Dossier inbox introuvable : {inbox_dir}")
        return

    files = [
        p for p in inbox_dir.iterdir()
        if p.is_file() and p.suffix.lower() in supported_extensions
    ]

    print(f"[INFO] Nombre de fichiers trouvés : {len(files)}")

    if not files:
        print(f"[INFO] Aucun fichier à traiter dans {inbox_dir}")
        return

    for file_path in files:
        try:
            process_file(file_path, extractor, ollama, config)
        except Exception as exc:
            print(f"[ERREUR] {file_path.name} : {exc}")
            traceback.print_exc()

    print("[INFO] Terminé.")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"[FATAL] {exc}")
        traceback.print_exc()